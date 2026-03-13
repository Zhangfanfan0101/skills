#!/usr/bin/env python3
"""
合同文件解析脚本

支持解析 Word (.docx) 和 PDF 文件，提取文本内容供智能体审查使用。
"""

import sys
import os
from pathlib import Path


def extract_docx(file_path: str) -> str:
    """解析 Word 文件（.docx格式）"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("缺少依赖包 python-docx，请运行: pip install python-docx==0.8.11")

    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)


def extract_pdf(file_path: str) -> str:
    """解析 PDF 文件"""
    try:
        import pdfplumber

        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)

                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            if row:
                                row_text = [str(cell).strip() if cell else "" for cell in row]
                                row_text = [t for t in row_text if t]
                                if row_text:
                                    text.append(" | ".join(row_text))

        return "\n".join(text)

    except ImportError:
        pass

    try:
        import PyPDF2

        text = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)

        return "\n".join(text)

    except ImportError:
        raise ImportError("缺少 PDF 解析依赖包，请安装: pip install pdfplumber==0.10.3")


def extract_contract(file_path: str) -> str:
    """自动识别文件类型并提取合同内容"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    file_ext = Path(file_path).suffix.lower()

    if file_ext == ".docx":
        return extract_docx(file_path)
    elif file_ext == ".pdf":
        return extract_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {file_ext}。仅支持 .docx 和 .pdf 格式")


def main():
    """命令行入口"""
    if len(sys.argv) < 2:
        print("使用方法: python scripts/extract_contract.py <合同文件路径>")
        print("支持的格式: .docx, .pdf")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        content = extract_contract(file_path)
        print(content)
    except Exception as e:
        print(f"解析失败: {str(e)}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
