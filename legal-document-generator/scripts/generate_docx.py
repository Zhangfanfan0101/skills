#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律文书生成器 - docx文档生成脚本

功能：将法律文书、相关法条、相关类案生成结构化的docx格式文档
修复问题：
1. 大段加粗不美观 - 仅对标题和关键词加粗，避免大段加粗
2. 包含\n字符 - 正确处理换行符
3. 多篇文书混在一起 - 正确识别和分隔多个文书
"""

import argparse
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_with_spacing(doc, text, level=1, font_size=None, font_name='仿宋'):
    """添加标题并调整间距"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    
    # 设置标题字体
    for run in heading.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size) if font_size else None
        run.font.color.rgb = None  # 确保字体颜色为黑色
        run.font.bold = True  # 标题默认加粗
    
    return heading


def clean_markdown(text):
    """
    清理markdown语法，智能处理加粗标记
    
    修复说明：
    - 移除所有markdown语法
    - 仅对标题和关键词加粗，避免大段加粗
    - 正确处理换行符\n
    """
    # 先统一换行符
    text = text.replace('\\n', '\n').replace('\\r', '\r')
    
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # 移除所有markdown加粗标记 **text** 和 __text__
        # 保留原始文本，不保留加粗标记
        cleaned_line = []
        i = 0
        while i < len(line):
            # 移除双星号（加粗）
            if i + 1 < len(line) and line[i] == '*' and line[i+1] == '*':
                i += 2
            # 移除双下划线（加粗）
            elif i + 1 < len(line) and line[i] == '_' and line[i+1] == '_':
                i += 2
            # 移除单星号（斜体）
            elif line[i] == '*':
                i += 1
            # 移除单下划线（斜体）
            elif line[i] == '_':
                i += 1
            else:
                cleaned_line.append(line[i])
                i += 1
        line = ''.join(cleaned_line)
        
        # 移除markdown标题标记 # ## ###
        line = re.sub(r'^#+\s*', '', line)
        
        # 移除markdown代码标记 `text`
        line = line.replace('`', '')
        
        # 移除markdown链接格式 [text](url) - 保留文本，移除链接
        line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        
        # 移除markdown列表标记 - * 1.，但保留列表内容
        line = line.strip()
        if line.startswith('- ') or line.startswith('* '):
            line = line[2:]
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            line = line[3:]
        
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)


def detect_multiple_documents(text):
    """
    检测文本中是否包含多个文档
    
    返回：
    - 如果检测到多个文档，返回文档列表
    - 否则返回包含原始文本的单元素列表
    
    检测规则：
    - 包含多个明显的文档标题（如"**民事起诉状**"、"**民事上诉状**"等）
    - 标题之间有明显的分隔（空行、分隔线等）
    """
    # 清理文本中的markdown标记以便检测
    cleaned = clean_markdown(text)
    
    # 查找可能的文档标题模式
    # 常见的法律文书标题模式
    doc_title_patterns = [
        r'^(.+状)$',
        r'^(.+书)$',
        r'^(.+申请书)$',
        r'^(.+意见书)$',
        r'^(.+声明书)$',
    ]
    
    lines = cleaned.split('\n')
    doc_indices = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        
        # 检查是否匹配文档标题模式
        for pattern in doc_title_patterns:
            if re.match(pattern, line):
                doc_indices.append(i)
                break
    
    # 如果找到多个文档标题
    if len(doc_indices) >= 2:
        documents = []
        for j in range(len(doc_indices)):
            start_idx = doc_indices[j]
            end_idx = doc_indices[j + 1] if j + 1 < len(doc_indices) else len(lines)
            doc_text = '\n'.join(lines[start_idx:end_idx]).strip()
            if doc_text:
                documents.append(doc_text)
        return documents
    
    # 否则返回原始文本
    return [cleaned]


def remove_duplicates(text):
    """去除重复的段落，保持顺序"""
    # 分割为非空行
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    # 去重并保持顺序
    seen = set()
    unique_lines = []
    for line in lines:
        if line not in seen:
            seen.add(line)
            unique_lines.append(line)
    return '\n'.join(unique_lines)


def add_section(doc, title, content, font_size=Pt(16), font_name='仿宋'):
    """
    添加文档部分
    
    修复说明：
    - 正确处理多文档情况，自动添加分隔线
    - 避免大段加粗，仅对特定关键词加粗
    - 正确处理换行符
    """
    # 添加部分标题（二级标题，使用三号字体=16磅）
    heading = doc.add_heading(title, level=2)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    
    # 设置二级标题字体：三号仿宋
    for run in heading.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = None  # 确保字体颜色为黑色
        run.font.bold = True  # 标题加粗

    # 清理markdown语法
    cleaned_content = clean_markdown(content)

    # 去除重复内容
    deduplicated_content = remove_duplicates(cleaned_content)

    # 检测是否包含多个文档
    documents = detect_multiple_documents(deduplicated_content)
    
    for doc_idx, doc_text in enumerate(documents):
        # 如果是第二个及以后的文档，添加分隔线
        if doc_idx > 0:
            sep_para = doc.add_paragraph('=' * 80)
            sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in sep_para.runs:
                run.font.color.rgb = None
            # 添加空行
            doc.add_paragraph()
        
        # 分割文档内容为段落
        paragraphs = doc_text.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # 判断是否为标题行（通常较短且在开头）
                # 标题行需要加粗，普通段落不加粗
                is_title_line = (
                    len(para_text) < 50 and
                    (para_text.endswith('状') or 
                     para_text.endswith('书') or
                     para_text.endswith('申请书') or
                     para_text.endswith('意见书') or
                     para_text.endswith('声明书'))
                )
                
                para = doc.add_paragraph(para_text)
                # 设置字体：三号仿宋
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = font_size
                    run.font.color.rgb = None  # 确保字体颜色为黑色
                    # 仅对标题行加粗，普通段落不加粗
                    run.font.bold = is_title_line
                
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_after = Pt(6)


def generate_legal_docx(legal_doc_content, laws_content, cases_content, output_path="legal-document.docx"):
    """
    生成法律文书docx文档

    参数：
        legal_doc_content: 法律文书正文内容
        laws_content: 相关法律法规内容
        cases_content: 相关类案内容
        output_path: 输出文件路径

    返回：
        生成的文件路径
    """
    try:
        # 创建文档对象
        doc = Document()

        # 设置默认字体：三号仿宋=16磅
        style = doc.styles['Normal']
        font = style.font
        font.name = '仿宋'
        font.size = Pt(16)  # 三号字体=16磅
        font.color.rgb = None  # 确保默认字体颜色为黑色

        # 添加文档标题：二号仿宋=22磅
        title = doc.add_heading('法律文书', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = '仿宋'
            run.font.size = Pt(22)  # 二号字体=22磅
            run.font.color.rgb = None  # 确保字体颜色为黑色
            run.font.bold = True  # 标题加粗

        # 添加分隔线
        para = doc.add_paragraph('_' * 80)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.color.rgb = None  # 确保分隔线颜色为黑色

        # 第一部分：法律文书正文
        add_section(doc, '一、法律文书正文', legal_doc_content)

        # 第二部分：相关法律法规
        if laws_content.strip():
            doc.add_page_break()
            add_section(doc, '二、相关法律法规', laws_content)

        # 第三部分：相关类案
        if cases_content.strip():
            doc.add_page_break()
            add_section(doc, '三、相关类案', cases_content)

        # 添加温馨提示
        doc.add_page_break()
        note_title = doc.add_heading('温馨提示', level=2)
        note_title.paragraph_format.space_before = Pt(12)
        note_title.paragraph_format.space_after = Pt(12)
        for run in note_title.runs:
            run.font.name = '仿宋'
            run.font.size = Pt(16)  # 三号字体
            run.font.color.rgb = None
            run.font.bold = True
        
        note_para = doc.add_paragraph(
            '本法律文书由AI智能体生成，仅供参考使用。重要法律事务建议咨询专业律师，以确保文书内容的准确性和法律效力。'
        )
        note_para.paragraph_format.line_spacing = 1.5
        # 确保温馨提示字体为三号仿宋，颜色为黑色
        for run in note_para.runs:
            run.font.name = '仿宋'
            run.font.size = Pt(16)  # 三号字体
            run.font.color.rgb = None

        # 保存文档
        doc.save(output_path)

        return output_path

    except Exception as e:
        raise Exception(f"生成docx文档失败: {str(e)}")


def parse_arguments():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(description='法律文书生成器 - 生成docx格式法律文书')

    parser.add_argument(
        '--legal-doc-content',
        type=str,
        required=True,
        help='法律文书正文内容'
    )

    parser.add_argument(
        '--laws-content',
        type=str,
        required=True,
        help='相关法律法规内容'
    )

    parser.add_argument(
        '--cases-content',
        type=str,
        required=True,
        help='相关类案内容'
    )

    parser.add_argument(
        '--output-path',
        type=str,
        default='legal-document.docx',
        help='输出文件路径（默认：legal-document.docx）'
    )

    return parser.parse_args()


def main():
    """主函数"""
    try:
        # 解析参数
        args = parse_arguments()

        # 验证输入
        if not args.legal_doc_content.strip():
            raise ValueError("法律文书正文内容不能为空")

        # 生成文档
        output_path = generate_legal_docx(
            args.legal_doc_content,
            args.laws_content,
            args.cases_content,
            args.output_path
        )

        print(f"✓ 法律文书已成功生成: {output_path}")
        return 0

    except ValueError as e:
        print(f"✗ 参数错误: {str(e)}", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"✗ 生成失败: {str(e)}", file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
